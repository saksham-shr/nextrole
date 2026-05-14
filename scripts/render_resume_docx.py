import io
import json
import sys

from docx import Document


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_bullets(doc, items):
    for item in items or []:
        doc.add_paragraph(str(item), style="List Bullet")


def main():
    raw = sys.stdin.read()
    data = json.loads(raw)

    doc = Document()

    name = data.get("name") or "Resume"
    doc.add_heading(name, 0)

    contact = data.get("contact") or {}
    contact_line = " | ".join(
        str(v) for v in [
            contact.get("email"),
            contact.get("phone"),
            contact.get("location"),
            contact.get("linkedin"),
            contact.get("github"),
            contact.get("portfolio"),
        ] if v
    )
    if contact_line:
      doc.add_paragraph(contact_line)

    summary = data.get("summary")
    if summary:
        add_heading(doc, "Summary")
        doc.add_paragraph(summary)

    competencies = data.get("competencies") or []
    if competencies:
        add_heading(doc, "Competencies")
        doc.add_paragraph(", ".join(str(x) for x in competencies))

    experience = data.get("experience") or []
    if experience:
        add_heading(doc, "Experience")
        for exp in experience:
            role = exp.get("role") or ""
            company = exp.get("company") or ""
            heading = " | ".join(part for part in [role, company] if part)
            if heading:
                p = doc.add_paragraph()
                r = p.add_run(heading)
                r.bold = True
            meta = " | ".join(part for part in [exp.get("period"), exp.get("location")] if part)
            if meta:
                doc.add_paragraph(meta)
            add_bullets(doc, exp.get("bullets") or [])

    projects = data.get("projects") or []
    if projects:
        add_heading(doc, "Projects")
        for project in projects:
            title = project.get("title") or ""
            description = project.get("description") or ""
            if title:
                p = doc.add_paragraph()
                r = p.add_run(title)
                r.bold = True
            if description:
                doc.add_paragraph(description)
            tech = project.get("tech") or []
            if tech:
                doc.add_paragraph("Tech: " + ", ".join(str(x) for x in tech))

    education = data.get("education") or []
    if education:
        add_heading(doc, "Education")
        for edu in education:
            doc.add_paragraph(" | ".join(str(v) for v in [edu.get("degree"), edu.get("institution"), edu.get("year")] if v))

    certifications = data.get("certifications") or []
    if certifications:
        add_heading(doc, "Certifications")
        for cert in certifications:
            doc.add_paragraph(" | ".join(str(v) for v in [cert.get("title"), cert.get("issuer"), cert.get("year")] if v))

    skills = data.get("skills") or {}
    if skills:
        add_heading(doc, "Skills")
        for category, items in skills.items():
            doc.add_paragraph(f"{category}: {', '.join(str(x) for x in items)}")

    output = io.BytesIO()
    doc.save(output)
    sys.stdout.buffer.write(output.getvalue())


if __name__ == "__main__":
    main()
